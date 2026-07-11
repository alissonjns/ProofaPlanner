import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


class Exportador:
    """Gerador de documentos Word para planos de aula e relatórios de alunos."""

    # Cores da identidade visual
    COR_PRIMARIA = RGBColor(0x10, 0xB9, 0x81)   # verde esmeralda
    COR_SECUNDARIA = RGBColor(0x05, 0x96, 0x69)  # verde escuro
    COR_TEXTO = RGBColor(0x1E, 0x29, 0x3B)       # slate escuro
    COR_ATINGIDO = RGBColor(0x05, 0x96, 0x69)
    COR_EM_DEV = RGBColor(0xD9, 0x77, 0x06)
    COR_NAO_INICIADO = RGBColor(0x94, 0xA3, 0xB8)

    def _configurar_documento(self) -> Document:
        """Cria e configura um documento Word com estilo base."""
        doc = Document()
        estilo = doc.styles["Normal"]
        estilo.font.name = "Calibri"
        estilo.font.size = Pt(11)
        return doc

    def _adicionar_cabecalho(self, doc: Document, titulo: str, subtitulo: str = ""):
        """Adiciona cabeçalho padronizado ao documento."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = self.COR_PRIMARIA

        if subtitulo:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(subtitulo)
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()

    def _adicionar_secao(self, doc: Document, titulo: str):
        """Adiciona título de seção com linha colorida."""
        p = doc.add_paragraph()
        run = p.add_run(f"▌ {titulo}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = self.COR_PRIMARIA

    def _salvar_bytes(self, doc: Document) -> bytes:
        """Salva o documento em memória e retorna os bytes."""
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ──────────────────────────────────────────────────────────────────────────
    # PLANO DE AULA
    # ──────────────────────────────────────────────────────────────────────────

    def gerar_plano_word(self, plano: dict) -> bytes:
        """Gera o documento Word do plano de aula (semanário)."""
        doc = self._configurar_documento()

        self._adicionar_cabecalho(
            doc,
            "PLANO DE AULA — EDUCAÇÃO INFANTIL",
            f"Gerado pelo ProfaPlanner em {datetime.now().strftime('%d/%m/%Y às %H:%M')}",
        )

        # ── Informações gerais ──────────────────────────────────────────────
        self._adicionar_secao(doc, "Informações Gerais")

        tabela = doc.add_table(rows=5, cols=2)
        tabela.style = "Table Grid"

        faixa = plano.get("faixa", "—")
        faixa_desc = ""
        if faixa == "EI01":
            faixa_desc = "EI01 — Bebês (0 a 1 ano e 6 meses)"
        elif faixa == "EI02":
            faixa_desc = "EI02 — Crianças bem pequenas (1a7m a 3a11m)"
        else:
            faixa_desc = "EI01 e EI02"

        linhas = [
            ("Tema / Atividade", plano.get("tema", "—")),
            ("Faixa Etária", faixa_desc),
            ("Campo de Experiência", plano.get("campo", "—")),
            ("Duração Estimada", plano.get("duracao", "—")),
            ("Data", datetime.now().strftime("%d/%m/%Y")),
        ]

        for i, (campo, valor) in enumerate(linhas):
            cell_campo = tabela.rows[i].cells[0]
            cell_valor = tabela.rows[i].cells[1]
            cell_campo.text = campo
            cell_valor.text = valor
            cell_campo.paragraphs[0].runs[0].bold = True
            cell_campo.paragraphs[0].runs[0].font.color.rgb = self.COR_SECUNDARIA

        doc.add_paragraph()

        # ── Objetivos BNCC ──────────────────────────────────────────────────
        self._adicionar_secao(doc, "Objetivos de Aprendizagem (BNCC)")

        objetivos = plano.get("objetivos", [])
        if objetivos:
            for obj in objetivos:
                p = doc.add_paragraph(style="List Bullet")
                run_codigo = p.add_run(f"[{obj['codigo']}]  ")
                run_codigo.bold = True
                run_codigo.font.color.rgb = self.COR_PRIMARIA
                p.add_run(obj["descricao"])
        else:
            doc.add_paragraph("Nenhum objetivo selecionado.")

        doc.add_paragraph()

        # ── Recursos ───────────────────────────────────────────────────────
        self._adicionar_secao(doc, "Recursos e Materiais")
        recursos = plano.get("recursos", "").strip()
        doc.add_paragraph(recursos if recursos else "A definir pela professora.")
        doc.add_paragraph()

        # ── Desenvolvimento ────────────────────────────────────────────────
        self._adicionar_secao(doc, "Desenvolvimento da Atividade")
        etapas = [
            "1. Acolhida e apresentação do tema:",
            "   _______________________________________________________________",
            "",
            "2. Desenvolvimento (descreva como a atividade será conduzida):",
            "   _______________________________________________________________",
            "   _______________________________________________________________",
            "",
            "3. Finalização e registro:",
            "   _______________________________________________________________",
        ]
        for linha in etapas:
            doc.add_paragraph(linha)

        doc.add_paragraph()

        # ── Avaliação ──────────────────────────────────────────────────────
        self._adicionar_secao(doc, "Registro e Avaliação")
        doc.add_paragraph(
            "Como a professora vai observar e registrar o desenvolvimento das crianças:"
        )
        doc.add_paragraph("   _______________________________________________________________")
        doc.add_paragraph("   _______________________________________________________________")

        doc.add_paragraph()
        doc.add_paragraph()

        # ── Assinatura ─────────────────────────────────────────────────────
        p_ass = doc.add_paragraph("_____________________________        _____________________________")
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_leg = doc.add_paragraph("    Professora Responsável                    Coordenação Pedagógica")
        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return self._salvar_bytes(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # RELATÓRIO BIMESTRAL DO ALUNO
    # ──────────────────────────────────────────────────────────────────────────

    def gerar_relatorio_aluno(self, aluno: dict, engine) -> bytes:
        """Gera o relatório bimestral de desenvolvimento do aluno."""
        doc = self._configurar_documento()

        self._adicionar_cabecalho(
            doc,
            "RELATÓRIO BIMESTRAL DE DESENVOLVIMENTO",
            "Educação Infantil — BNCC",
        )

        # ── Dados da criança ────────────────────────────────────────────────
        self._adicionar_secao(doc, "Dados da Criança")

        faixa = aluno.get("faixa", "EI02")
        faixa_desc = (
            "EI01 — Bebês (0 a 1 ano e 6 meses)"
            if faixa == "EI01"
            else "EI02 — Crianças bem pequenas (1a7m a 3a11m)"
        )

        tabela_dados = doc.add_table(rows=3, cols=2)
        tabela_dados.style = "Table Grid"
        linhas_dados = [
            ("Nome da Criança", aluno.get("nome", "—")),
            ("Faixa Etária", faixa_desc),
            ("Período de Referência", datetime.now().strftime("%B de %Y").capitalize()),
        ]
        for i, (campo, valor) in enumerate(linhas_dados):
            tabela_dados.rows[i].cells[0].text = campo
            tabela_dados.rows[i].cells[1].text = valor
            tabela_dados.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            tabela_dados.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = self.COR_SECUNDARIA

        doc.add_paragraph()

        # ── Evolução por campo ──────────────────────────────────────────────
        self._adicionar_secao(doc, "Desenvolvimento por Campo de Experiência")

        status_map = {
            "A": ("✅", "Atingido"),
            "D": ("🔄", "Em Desenvolvimento"),
            "N": ("⏳", "Não Iniciado"),
        }
        cores_status = {
            "A": self.COR_ATINGIDO,
            "D": self.COR_EM_DEV,
            "N": self.COR_NAO_INICIADO,
        }

        evolucao = aluno.get("evolucao", {})
        objetivos_faixa = engine.get_objetivos_por_faixa(faixa)

        for campo_nome, objetivos in objetivos_faixa.items():
            # Sub-título do campo
            p_campo = doc.add_paragraph()
            run_campo = p_campo.add_run(f"\n{campo_nome}")
            run_campo.bold = True
            run_campo.font.size = Pt(11)
            run_campo.font.color.rgb = self.COR_TEXTO

            for obj in objetivos:
                status_codigo = evolucao.get(obj["codigo"], "N")
                emoji, status_texto = status_map[status_codigo]
                cor = cores_status[status_codigo]

                p = doc.add_paragraph(style="List Bullet")
                run_cod = p.add_run(f"[{obj['codigo']}] ")
                run_cod.bold = True
                run_cod.font.color.rgb = self.COR_PRIMARIA
                p.add_run(f"{obj['descricao']}  ")
                run_st = p.add_run(f"— {emoji} {status_texto}")
                run_st.italic = True
                run_st.font.color.rgb = cor

        doc.add_paragraph()

        # ── Síntese narrativa ───────────────────────────────────────────────
        self._adicionar_secao(doc, "Síntese do Período")

        total = len(evolucao)
        atingidos = sum(1 for v in evolucao.values() if v == "A")
        em_dev = sum(1 for v in evolucao.values() if v == "D")
        nome = aluno.get("nome", "A criança")

        sintese = (
            f"{nome} apresentou desenvolvimento em {total} objetivo(s) de aprendizagem "
            f"avaliados neste período. Foram {atingidos} objetivo(s) atingido(s) e "
            f"{em_dev} em processo de desenvolvimento. "
            f"A criança demonstra evolução contínua em seu percurso de aprendizagem, "
            f"conforme as observações e registros realizados pela professora ao longo do bimestre."
        )
        doc.add_paragraph(sintese)

        doc.add_paragraph()

        # ── Observações complementares ──────────────────────────────────────
        self._adicionar_secao(doc, "Observações Complementares da Professora")
        doc.add_paragraph("   _______________________________________________________________")
        doc.add_paragraph("   _______________________________________________________________")
        doc.add_paragraph("   _______________________________________________________________")

        doc.add_paragraph()
        doc.add_paragraph()

        # ── Assinatura ─────────────────────────────────────────────────────
        p_ass = doc.add_paragraph("_____________________________        _____________________________")
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_leg = doc.add_paragraph("    Professora Responsável                    Coordenação Pedagógica")
        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return self._salvar_bytes(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SEQUÊNCIA DIDÁTICA — WORD
    # ──────────────────────────────────────────────────────────────────────────

    def gerar_sd_word(self, dados: dict) -> bytes:
        """Gera Sequência Didática em Word, no formato real usado pela professora."""
        doc = self._configurar_documento()

        # Título principal
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_title.add_run("Sequência Didática")
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = self.COR_PRIMARIA

        # Subtítulo (nome da SD)
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p_sub.add_run(dados.get("titulo", ""))
        r2.bold = True
        r2.font.size = Pt(16)
        r2.font.color.rgb = self.COR_TEXTO

        doc.add_paragraph()

        # Tabela de informações gerais
        infos = [
            ("Tema",                  dados.get("tema", "—")),
            ("Duração",               dados.get("duracao", "—")),
            ("Campo de Experiência",  dados.get("campo", "—")),
            ("Segmento / Turma",      dados.get("segmento", "—")),
            ("Faixa Etária",          dados.get("faixa", "—")),
        ]
        tbl = doc.add_table(rows=len(infos), cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate(infos):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = v
            runs = tbl.rows[i].cells[0].paragraphs[0].runs
            if runs:
                runs[0].bold = True
                runs[0].font.color.rgb = self.COR_SECUNDARIA

        doc.add_paragraph()

        # Justificativa
        self._adicionar_secao(doc, "Justificativa")
        j = dados.get("justificativa", "").strip()
        doc.add_paragraph(j if j else "_" * 80)
        doc.add_paragraph()

        # Objetivo Geral
        self._adicionar_secao(doc, "Objetivo Geral")
        og = dados.get("obj_geral", "").strip()
        doc.add_paragraph(og if og else "_" * 80)
        doc.add_paragraph()

        # Objetivos BNCC
        self._adicionar_secao(doc, "Objetivos de Aprendizagem (BNCC)")
        for obj in dados.get("objetivos_bncc", []):
            p = doc.add_paragraph()
            rc = p.add_run(f"{obj['codigo']}— ")
            rc.bold = True
            rc.font.color.rgb = self.COR_PRIMARIA
            p.add_run(obj["descricao"])
        doc.add_paragraph()

        # Atividades numeradas
        atividades = [a for a in dados.get("atividades", [])
                      if a.get("nome") or a.get("descricao")]
        for i, ativ in enumerate(atividades, 1):
            p_ativ = doc.add_paragraph()
            ra = p_ativ.add_run(f"Atividade {i}: {ativ.get('nome', '')}")
            ra.bold = True
            ra.font.size = Pt(11)
            ra.font.color.rgb = self.COR_TEXTO
            desc = ativ.get("descricao", "").strip()
            doc.add_paragraph(desc if desc else "_" * 80)
            doc.add_paragraph()

        # Avaliação
        self._adicionar_secao(doc, "Avaliação")
        av = dados.get("avaliacao", "").strip()
        doc.add_paragraph(av if av else "_" * 80)

        doc.add_paragraph()
        doc.add_paragraph()
        p_ass = doc.add_paragraph(
            "_____________________________        _____________________________"
        )
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_leg = doc.add_paragraph(
            "    Professora Responsável                    Coordenação Pedagógica"
        )
        p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return self._salvar_bytes(doc)

    # ──────────────────────────────────────────────────────────────────────────
    # SEQUÊNCIA DIDÁTICA — PDF
    # ──────────────────────────────────────────────────────────────────────────

    def gerar_sd_pdf(self, dados: dict) -> bytes:
        """Gera Sequência Didática em PDF profissional usando ReportLab."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
        )

        buffer = io.BytesIO()
        doc_pdf = SimpleDocTemplate(
            buffer, pagesize=A4,
            leftMargin=2.5 * cm, rightMargin=2.5 * cm,
            topMargin=2 * cm, bottomMargin=2.5 * cm,
            title=f"SD - {dados.get('titulo', '')}",
        )

        cor_verde       = colors.HexColor("#10B981")
        cor_verde_esc   = colors.HexColor("#059669")
        cor_verde_claro = colors.HexColor("#F0FDF4")
        cor_verde_borda = colors.HexColor("#D1FAE5")
        cor_texto       = colors.HexColor("#1E293B")
        cor_cinza       = colors.HexColor("#64748B")

        s = getSampleStyleSheet()
        s_titulo   = ParagraphStyle("s_t",  fontSize=22, textColor=cor_verde,
                                    fontName="Helvetica-Bold", alignment=1, spaceAfter=4)
        s_sub      = ParagraphStyle("s_s",  fontSize=16, textColor=cor_texto,
                                    fontName="Helvetica-Bold", alignment=1, spaceAfter=14)
        s_h2       = ParagraphStyle("s_h2", fontSize=12, textColor=cor_verde,
                                    fontName="Helvetica-Bold", spaceBefore=14, spaceAfter=6)
        s_h3       = ParagraphStyle("s_h3", fontSize=11, textColor=cor_texto,
                                    fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4)
        s_body     = ParagraphStyle("s_b",  fontSize=10, textColor=cor_texto,
                                    leading=16, spaceAfter=6)
        s_vazio    = ParagraphStyle("s_v",  fontSize=9,  textColor=cor_cinza,
                                    leading=14, spaceAfter=4)

        story = []

        # Cabeçalho
        story.append(Paragraph("Sequência Didática", s_titulo))
        story.append(Paragraph(dados.get("titulo", ""), s_sub))
        story.append(HRFlowable(width="100%", color=cor_verde, thickness=2, spaceAfter=10))

        # Tabela de informações
        linhas = [
            ["Tema",                 dados.get("tema", "—")],
            ["Duração",              dados.get("duracao", "—")],
            ["Campo de Experiência", dados.get("campo", "—")],
            ["Segmento / Turma",     dados.get("segmento", "—")],
            ["Faixa Etária",         dados.get("faixa", "—")],
        ]
        tbl = Table(linhas, colWidths=[5 * cm, 11.5 * cm])
        tbl.setStyle(TableStyle([
            ("FONTNAME",   (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME",   (1, 0), (1, -1), "Helvetica"),
            ("FONTSIZE",   (0, 0), (-1, -1), 10),
            ("TEXTCOLOR",  (0, 0), (0, -1), cor_verde_esc),
            ("TEXTCOLOR",  (1, 0), (1, -1), cor_texto),
            ("BACKGROUND", (0, 0), (0, -1), cor_verde_claro),
            ("GRID",       (0, 0), (-1, -1), 0.5, cor_verde_borda),
            ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING",  (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING",   (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 7),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.4 * cm))

        def secao(titulo: str, texto: str):
            story.append(HRFlowable(width="100%", color=cor_verde_borda, thickness=0.5))
            story.append(Paragraph(f"▌ {titulo}", s_h2))
            t = texto.strip() if texto else ""
            story.append(Paragraph(
                t.replace("\n", "<br/>") if t else "___________________________________________",
                s_body if t else s_vazio,
            ))
            story.append(Spacer(1, 0.15 * cm))

        secao("Justificativa", dados.get("justificativa", ""))
        secao("Objetivo Geral", dados.get("obj_geral", ""))

        # Objetivos BNCC
        story.append(HRFlowable(width="100%", color=cor_verde_borda, thickness=0.5))
        story.append(Paragraph("▌ Objetivos de Aprendizagem (BNCC)", s_h2))
        for obj in dados.get("objetivos_bncc", []):
            story.append(Paragraph(
                f'<font color="#059669" name="Courier"><b>[{obj["codigo"]}]</b></font>'
                f'&nbsp;&nbsp;{obj["descricao"]}',
                s_body,
            ))
        story.append(Spacer(1, 0.2 * cm))

        # Atividades
        atividades = [a for a in dados.get("atividades", [])
                      if a.get("nome") or a.get("descricao")]
        if atividades:
            story.append(HRFlowable(width="100%", color=cor_verde_borda, thickness=0.5))
            story.append(Paragraph("▌ Atividades", s_h2))
            for i, ativ in enumerate(atividades, 1):
                story.append(Paragraph(f"Atividade {i}: {ativ.get('nome', '')}", s_h3))
                desc = ativ.get("descricao", "").strip()
                story.append(Paragraph(
                    desc.replace("\n", "<br/>") if desc else "___________________________________________",
                    s_body if desc else s_vazio,
                ))
                story.append(Spacer(1, 0.1 * cm))

        secao("Avaliação", dados.get("avaliacao", ""))

        # Assinatura
        story.append(Spacer(1, 1.2 * cm))
        story.append(HRFlowable(width="100%", color=cor_verde_borda, thickness=0.5))
        story.append(Spacer(1, 1.5 * cm))
        ass = Table(
            [["_____________________________", "_____________________________"],
             ["Professora Responsável", "Coordenação Pedagógica"]],
            colWidths=[8 * cm, 8 * cm],
        )
        ass.setStyle(TableStyle([
            ("FONTNAME",  (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE",  (0, 0), (-1, -1), 9),
            ("TEXTCOLOR", (0, 0), (-1, -1), cor_cinza),
            ("ALIGN",     (0, 0), (-1, -1), "CENTER"),
            ("TOPPADDING", (0, 1), (-1, 1), 2),
        ]))
        story.append(ass)

        doc_pdf.build(story)
        buffer.seek(0)
        return buffer.getvalue()

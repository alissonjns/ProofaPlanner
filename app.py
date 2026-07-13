"""
ProfaPlanner — Assistente de Planejamento Pedagógico BNCC
Educação Infantil · EI01 e EI02
"""

import io
import random
from pathlib import Path

import streamlit as st

from modules.bncc_engine import BNCCEngine
from modules.chatbot import ProfaBot
from modules.exportador import Exportador

# ─── Configuração da Página ───────────────────────────────────────────────────
st.set_page_config(
    page_title="ProfaPlanner | Seu assistente de planejamento",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CSS ─────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

/* ── Reset & Base ── */
html, body { font-family: 'Inter', sans-serif; }
html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"] {
    background-color: #F8FAFC !important;
    color: #1E293B !important;
}
.block-container { 
    padding-top: 2rem; padding-bottom: 3rem; max-width: 960px; 
    background-color: #F8FAFC;
}

/* ── Esconde elementos do Streamlit Cloud ── */
#MainMenu {visibility: hidden;}
header {display: none !important;}
footer {display: none !important;}
.viewerBadge_container__1QSob {display: none !important;}
[data-testid="manage-app-button"] {display: none !important;}

/* ── Fundo geral forçado ── */
[data-testid="stMain"] { background-color: #F8FAFC !important; }

/* ── Divider ── */
hr { border-color: #E2E8F0 !important; }

/* ── Inputs e Placeholders ── */
::placeholder {
    color: #94A3B8 !important;
    opacity: 1 !important;
}
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea {
    background-color: #FFFFFF !important;
    border: 1.5px solid #E2E8F0 !important;
    border-radius: 8px !important;
    color: #1E293B !important;
    font-size: 0.92rem !important;
}
[data-testid="stTextInput"] input:focus,
[data-testid="stTextArea"] textarea:focus {
    border-color: #10B981 !important;
    box-shadow: 0 0 0 3px rgba(16,185,129,0.12) !important;
}

/* ── Selectbox ── */
[data-testid="stSelectbox"] > div > div,
[data-testid="stSelectbox"] > div > div > div,
[data-testid="stSelectbox"] input,
[data-baseweb="select"] > div,
[data-baseweb="select"] > div > div,
[data-baseweb="popover"] li,
[data-baseweb="popover"] ul {
    background-color: #FFFFFF !important;
    border-color: #E2E8F0 !important;
    border-radius: 8px !important;
    color: #1E293B !important;
}
/* Força o texto do item selecionado a ser visível */
[data-baseweb="select"] span,
[data-baseweb="select"] div[class*="singleValue"],
[class*="placeholder"],
[class*="singleValue"] {
    color: #1E293B !important;
}
/* Dropdown aberto */
[data-baseweb="popover"] {
    background: #FFFFFF !important;
}
[data-baseweb="menu"] {
    background: #FFFFFF !important;
}
[data-baseweb="menu"] li {
    color: #1E293B !important;
    background: #FFFFFF !important;
}
[data-baseweb="menu"] li:hover {
    background: #F0FDF4 !important;
    color: #059669 !important;
}

/* ── Labels ── */
label, [data-testid="stWidgetLabel"] {
    color: #475569 !important;
    font-weight: 500 !important;
    font-size: 0.88rem !important;
}

/* ── Primary buttons ── */
[data-testid="stButton"] button[kind="primary"] {
    background: linear-gradient(135deg, #10B981, #059669) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    font-size: 0.95rem !important;
    padding: 0.65rem 1.5rem !important;
    box-shadow: 0 4px 15px rgba(16,185,129,0.3) !important;
    transition: all 0.2s ease !important;
}
[data-testid="stButton"] button[kind="primary"]:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(16,185,129,0.4) !important;
}

/* ── Secondary buttons ── */
[data-testid="stButton"] button[kind="secondary"] {
    background: #FFFFFF !important;
    color: #475569 !important;
    border: 1.5px solid #E2E8F0 !important;
    border-radius: 10px !important;
    font-weight: 500 !important;
}
[data-testid="stButton"] button[kind="secondary"]:hover {
    border-color: #10B981 !important;
    color: #10B981 !important;
}

/* ── Download buttons ── */
[data-testid="stDownloadButton"] button {
    border-radius: 10px !important;
    font-weight: 600 !important;
    color: #1E293B !important;
    background: #FFFFFF !important;
    border: 1.5px solid #E2E8F0 !important;
}
[data-testid="stDownloadButton"] button[kind="primary"] {
    background: linear-gradient(135deg, #10B981, #059669) !important;
    color: white !important;
    border-color: transparent !important;
}

/* ── Chat messages ── */
[data-testid="stChatMessage"] {
    background: #FFFFFF !important;
    border: 1px solid #E2E8F0 !important;
    border-radius: 12px !important;
    color: #1E293B !important;
}
[data-testid="stChatMessage"] * {
    color: #1E293B !important;
}
[data-testid="stChatMessage"] p,
[data-testid="stChatMessage"] span,
[data-testid="stChatMessage"] div {
    color: #1E293B !important;
}
/* Mensagem do usuário */
[data-testid="stChatMessage"][data-testid*="user"],
[data-testid="stChatMessageContent"] {
    color: #1E293B !important;
}

/* ── Chat input ── */
[data-testid="stChatInput"],
[data-testid="stChatInput"] > div {
    background-color: #FFFFFF !important;
    border: 1.5px solid #E2E8F0 !important;
    border-radius: 12px !important;
}
[data-testid="stChatInput"] textarea {
    background-color: #FFFFFF !important;
    color: #1E293B !important;
    border: none !important;
}
[data-testid="stChatInput"] textarea::placeholder {
    color: #94A3B8 !important;
}
[data-testid="stChatInput"] button {
    background: linear-gradient(135deg, #10B981, #059669) !important;
    border-radius: 8px !important;
}

/* ── Paragraphs e textos gerais ── */
p { color: #1E293B !important; }
.stMarkdown { color: #1E293B !important; }
.stMarkdown p { color: #1E293B !important; }

/* ── Checkboxes ── */
[data-testid="stCheckbox"] { color: #1E293B !important; }
[data-testid="stCheckbox"] label { color: #334155 !important; font-size: 0.88rem !important; }

/* ── Expanders ── */
[data-testid="stExpander"] {
    background-color: #FFFFFF !important;
    border: 1.5px solid #E2E8F0 !important;
    border-radius: 12px !important;
    margin-bottom: 0.5rem !important;
}
[data-testid="stExpander"] summary {
    color: #1E293B !important;
    font-weight: 600 !important;
    background-color: #FFFFFF !important;
}
[data-testid="stExpanderDetails"] {
    background-color: #FFFFFF !important;
    color: #1E293B !important;
}

/* ── File Uploader ── */
[data-testid="stFileUploaderDropzone"] {
    background-color: #F8FAFC !important;
    border: 2px dashed #CBD5E1 !important;
    border-radius: 12px !important;
    padding: 1.5rem !important;
}
/* Esconder o botãozinho interno para não parecer "duas caixas" */
[data-testid="stFileUploaderDropzone"] button {
    display: none !important;
}
[data-testid="stFileUploader"] {
    color: #1E293B !important;
}

/* ── Info / warning / success ── */
[data-testid="stAlert"] {
    border-radius: 10px !important;
    border-left-width: 4px !important;
}

/* ── Tabs ── */
[data-testid="stTabs"] [data-baseweb="tab-list"] {
    background: #F1F5F9 !important;
    border-radius: 10px !important;
    padding: 4px !important;
}
[data-testid="stTabs"] [data-baseweb="tab"] {
    color: #64748B !important;
    border-radius: 8px !important;
    font-weight: 500 !important;
}
[data-testid="stTabs"] [aria-selected="true"] {
    background-color: #FFFFFF !important;
    color: #10B981 !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.08) !important;
}

/* ── Sidebar ── */
[data-testid="stSidebar"] { 
    background-color: #FFFFFF !important;
    border-right: 1px solid #E2E8F0 !important;
}

/* ── Pills navigation ── */
[data-testid="stPills"] {
    gap: 6px !important;
}
[data-testid="stPills"] button {
    background: #FFFFFF !important;
    color: #475569 !important;
    border: 1.5px solid #E2E8F0 !important;
    border-radius: 20px !important;
    font-weight: 500 !important;
    font-size: 0.85rem !important;
    transition: all 0.2s ease !important;
}
[data-testid="stPills"] button:hover {
    border-color: #10B981 !important;
    color: #10B981 !important;
}
[data-testid="stPills"] button[aria-selected="true"] {
    background: linear-gradient(135deg, #10B981, #059669) !important;
    color: white !important;
    border-color: transparent !important;
    box-shadow: 0 4px 12px rgba(16,185,129,0.3) !important;
}

/* ── Spinner ── */
[data-testid="stSpinner"] { color: #10B981 !important; }

/* ── Progress bar ── */
.prog-bar-bg { background:#E2E8F0; border-radius:99px; height:6px; overflow:hidden; }
.prog-bar { background: linear-gradient(90deg, #10B981, #34D399); height:100%; border-radius:99px; transition:width 0.4s ease; }

/* ── Custom cards ── */
.pp-card {
    background: #FFFFFF;
    border: 1.5px solid #E2E8F0;
    border-radius: 14px;
    padding: 1.5rem;
    margin-bottom: 1rem;
    box-shadow: 0 2px 12px rgba(0,0,0,0.05);
}
.pp-card-verde {
    background: linear-gradient(135deg, #ECFDF5, #F0FDF4);
    border: 1.5px solid #A7F3D0;
    border-radius: 14px;
    padding: 1.5rem;
    margin-bottom: 1rem;
}
.pp-card-coral {
    background: linear-gradient(135deg, #FFF7ED, #FFEDD5);
    border: 1.5px solid #FED7AA;
    border-radius: 14px;
    padding: 1.5rem;
    margin-bottom: 1rem;
}
.pp-card-azul {
    background: linear-gradient(135deg, #EFF6FF, #DBEAFE);
    border: 1.5px solid #BFDBFE;
    border-radius: 14px;
    padding: 1.5rem;
    margin-bottom: 1rem;
}
.pp-step {
    display: flex;
    align-items: center;
    gap: 12px;
    padding: 0.8rem 1.2rem;
    border-radius: 99px;
    font-size: 0.88rem;
    font-weight: 600;
}
.pp-step-active {
    background: linear-gradient(135deg, #10B981, #059669);
    color: white;
    box-shadow: 0 4px 15px rgba(16,185,129,0.3);
}
.pp-step-done {
    background: #ECFDF5;
    color: #059669;
    border: 1.5px solid #A7F3D0;
}
.pp-step-pending {
    background: #F8FAFC;
    color: #94A3B8;
    border: 1.5px solid #E2E8F0;
}
.pp-step-num {
    width: 26px; height: 26px;
    border-radius: 50%;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    font-size: 0.8rem;
    font-weight: 700;
}
.pp-hero-title {
    font-size: 2.2rem;
    font-weight: 800;
    color: #0F172A;
    line-height: 1.2;
}
.pp-hero-sub {
    font-size: 1rem;
    color: #64748B;
    line-height: 1.6;
}
.pp-section-title {
    font-size: 1.1rem;
    font-weight: 700;
    color: #0F172A;
    margin-bottom: 0.5rem;
}
.pp-tag {
    display: inline-block;
    padding: 3px 10px;
    border-radius: 99px;
    font-size: 0.78rem;
    font-weight: 600;
}
.pp-tag-verde { background: #ECFDF5; color: #059669; }
.pp-tag-coral { background: #FFF7ED; color: #EA580C; }
.pp-tag-azul  { background: #EFF6FF; color: #2563EB; }
.pp-obj-card {
    background: #FFFFFF;
    border: 1.5px solid #E2E8F0;
    border-radius: 10px;
    padding: 0.8rem 1rem;
    margin-bottom: 0.5rem;
    transition: border-color 0.2s ease;
}
.pp-obj-card:hover { border-color: #10B981; }
.pp-divider { height: 1px; background: #E2E8F0; margin: 1.5rem 0; }

/* ── Cabeçalho / navegação ── */
.pp-navbar {
    background: #FFFFFF;
    border-bottom: 1px solid #E2E8F0;
    padding: 1rem 0 1.5rem;
    margin-bottom: 1.5rem;
}
/* ── Seção título (classe antiga) ── */
.sec {
    font-size: 1.5rem;
    font-weight: 800;
    color: #0F172A;
    margin-bottom: 0.3rem;
}
/* ── Corrigir texto padrão do Streamlit ── */
[data-testid="stMarkdown"] * { color: #1E293B; }
[data-testid="stText"] { color: #1E293B !important; }
/* Corrige tag-ei01 e tag-ei02 */
.tag-ei01 { background:#DBEAFE; color:#1D4ED8; padding:2px 8px; border-radius:99px; font-size:0.78rem; font-weight:600; }
.tag-ei02 { background:#FEF3C7; color:#D97706; padding:2px 8px; border-radius:99px; font-size:0.78rem; font-weight:600; }
</style>
<meta name="google" content="notranslate">
""", unsafe_allow_html=True)

# ─── Session State ────────────────────────────────────────────────────────────
_def = {
    "pagina": "boas_vindas",   # começa no tour obrigatório
    "tour_step": 1,
    "alunos": [],
    "historico_chat": [],
    "modelo_escola": None,
    "plano_atual": {},
    "gemini_key": "",
}
for _k, _v in _def.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ─── Motor BNCC ───────────────────────────────────────────────────────────────
@st.cache_resource(show_spinner="Carregando base BNCC...")
def _engine():
    return BNCCEngine()
engine = _engine()


# ─── Detecta documentos na pasta do projeto ───────────────────────────────────
def detectar_arquivos_locais() -> list[Path]:
    """Retorna .docx e .xlsx encontrados na mesma pasta do app.py"""
    pasta = Path(__file__).parent
    return sorted(
        [f for f in pasta.iterdir()
         if f.suffix.lower() in (".docx", ".xlsx", ".xls", ".doc")],
        key=lambda f: f.name
    )


# ─── Top Navigation ─────────────────────────────────────────────────────────────
_NA_TOUR = st.session_state.pagina == "boas_vindas"

if not _NA_TOUR:
    st.markdown("""
    <div class='pp-navbar' style='display:flex; align-items:center; gap:14px;'>
        <div style='width:46px; height:46px; background:linear-gradient(135deg,#10B981,#059669); border-radius:12px; display:flex; align-items:center; justify-content:center; font-size:1.4rem; box-shadow:0 4px 12px rgba(16,185,129,0.3);'>📚</div>
        <div>
            <div style='font-size:1.5rem; font-weight:800; color:#0F172A; line-height:1;'>ProfaPlanner</div>
            <div style='font-size:0.72rem; color:#10B981; letter-spacing:1.5px; font-weight:700; text-transform:uppercase;'>BNCC · Educação Infantil</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    _nav_options = {
        "Inicio": ("inicio", "🏠", "Início"),
        "Modelo": ("modelo", "📋", "Meu Modelo"),
        "Plano":  ("plano",  "✏️", "Criar Plano"),
        "Alunos": ("alunos", "👶", "Alunos"),
        "Bot":    ("profabot","💬", "ProfaBot"),
    }
    nav_cols = st.columns(len(_nav_options))
    for col, (key, (pag, icon, label)) in zip(nav_cols, _nav_options.items()):
        with col:
            is_active = st.session_state.pagina == pag
            if st.button(f"{icon} {label}", key=f"nav_{key}", use_container_width=True,
                         type="primary" if is_active else "secondary"):
                st.session_state.pagina = pag
                st.rerun()

    st.divider()


# ═══════════════════════════════════════════════════════════════════════════════
#  TOUR OBRIGATÓRIO DE BOAS-VINDAS
# ═══════════════════════════════════════════════════════════════════════════════
def pagina_boas_vindas():
    # Esconde sidebar durante o tour
    st.markdown("""<style>[data-testid="stSidebar"]{display:none!important}</style>""",
                unsafe_allow_html=True)

    step = st.session_state.tour_step

    # ── Barra de progresso ────────────────────────────────────────────────
    pct = int((step / 3) * 100)
    st.markdown(f"""
    <div style="max-width:700px; margin:0 auto 2rem;">
        <div style="display:flex; justify-content:space-between; color:#475569; font-size:0.78rem; margin-bottom:6px;">
            <span>Apresentação do sistema</span>
            <span>Passo {step} de 3</span>
        </div>
        <div class="prog-bar-bg"><div class="prog-bar" style="width:{pct}%"></div></div>
    </div>
    """, unsafe_allow_html=True)

    col_c, _ = st.columns([2, 1])  # centraliza conteúdo
    with col_c:

        # ── PASSO 1: Boas-vindas ──────────────────────────────────────────
        if step == 1:
            st.markdown("""
            <div style="text-align:center; padding:2rem 0 2rem;">
                <div style="font-size:4.5rem; margin-bottom:1rem;">📚</div>
                <div style="font-size:2.8rem; font-weight:800; color:#0F172A; line-height:1.2; margin-bottom:0.8rem;">
                    Olá, Professora!
                </div>
                <div style="font-size:1.05rem; color:#475569; max-width:500px; margin:0 auto; line-height:1.7;">
                    O <strong style="color:#10B981;">ProfaPlanner</strong> foi criado para facilitar a sua vida.<br>
                    Chega de perder tempo procurando códigos da BNCC — o sistema faz isso por você! 🎉
                </div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("""
            <div class="pp-card-verde" style="text-align:center; padding:1.8rem;">
                <div style="color:#059669; font-weight:700; font-size:1.05rem; margin-bottom:1rem;">
                    ✨ O que você vai conseguir fazer aqui:
                </div>
                <div style="color:#065F46; font-size:0.97rem; line-height:2.4;">
                    📝 &nbsp;<strong>Gerar planos de aula</strong> com os objetivos BNCC corretos em minutos<br>
                    👶 &nbsp;<strong>Criar o relatório bimestral</strong> de cada aluno com facilidade<br>
                    💬 &nbsp;<strong>Tirar dúvidas</strong> sobre a BNCC com uma assistente virtual
                </div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Próximo →  Entender como funciona", type="primary", use_container_width=True):
                st.session_state.tour_step = 2
                st.rerun()

        # ── PASSO 2: Como funciona ────────────────────────────────────────
        elif step == 2:
            st.markdown("""
            <div style="font-size:2rem; font-weight:800; color:#0F172A; text-align:center; margin-bottom:1.5rem;">
                Como o sistema funciona?
            </div>
            """, unsafe_allow_html=True)

            _passos = [
                ("📋", "#EFF6FF", "#1D4ED8", "#BFDBFE", "1. Você coloca o modelo da escola",
                 "Se a escola te deu um formulário de plano de aula no Word ou Excel, você o adiciona aqui. O sistema usa ele como base. Se não tiver, usamos um modelo pronto."),
                ("✏️", "#ECFDF5", "#065F46", "#A7F3D0", "2. Você digita o tema da aula",
                 "Escreve o tema que você planejou, como 'Brincadeira com argila' ou 'Música e ritmo'. O sistema encontra automaticamente os objetivos BNCC certos para você!"),
                ("⬇️", "#FFF7ED", "#9A3412", "#FED7AA", "3. Você baixa o documento pronto",
                 "Clique em um botão e o plano de aula já sai preenchido no formato Word, pronto para imprimir ou entregar à coordenação."),
            ]
            for icon, bg, cor, border, titulo, desc in _passos:
                st.markdown(f"""
                <div style="background:{bg}; border:1.5px solid {border}; border-radius:14px; padding:1.3rem 1.5rem; margin-bottom:0.8rem;">
                    <div style="font-size:1.5rem; margin-bottom:0.4rem;">{icon}</div>
                    <div style="font-weight:700; color:{cor}; margin-bottom:0.3rem; font-size:1rem;">{titulo}</div>
                    <div style="color:#475569; font-size:0.88rem; line-height:1.6;">{desc}</div>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            c1, c2 = st.columns(2)
            with c1:
                if st.button("← Voltar", use_container_width=True):
                    st.session_state.tour_step = 1
                    st.rerun()
            with c2:
                if st.button("Próximo →  Começar agora", type="primary", use_container_width=True):
                    st.session_state.tour_step = 3
                    st.rerun()

        # ── PASSO 3: Começar ──────────────────────────────────────────────
        elif step == 3:
            st.markdown("""
            <div style="font-size:2rem; font-weight:800; color:#0F172A; text-align:center; margin-bottom:0.5rem;">Tudo pronto! 🎉</div>
            <div style="color:#64748B; font-size:1rem; text-align:center; margin-bottom:1.5rem;">Por onde você quer começar?</div>
            """, unsafe_allow_html=True)

            _opcoes = [
                ("modelo",   "📋", "Carregar o modelo da escola",
                 "Tenho o arquivo que a escola me deu e quero carregá-lo"),
                ("plano",    "✏️", "Criar meu primeiro plano de aula",
                 "Quero digitar o tema de uma aula e gerar o documento agora"),
                ("profabot", "💬", "Fazer uma pergunta sobre a BNCC",
                 "Tenho uma dúvida e quero perguntar para a assistente virtual"),
            ]
            for pag, icon, titulo, desc in _opcoes:
                st.markdown(f"""
                <div style="background:#FFFFFF; border:1.5px solid #E2E8F0; border-radius:14px; padding:1rem 1.3rem; display:flex; align-items:center; gap:1rem; margin-bottom:0.4rem; box-shadow:0 2px 8px rgba(0,0,0,0.04);">
                    <div style="font-size:2rem; flex-shrink:0;">{icon}</div>
                    <div>
                        <div style="font-weight:700; color:#0F172A; font-size:0.95rem;">{titulo}</div>
                        <div style="color:#64748B; font-size:0.83rem;">{desc}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                if st.button(f"Ir para: {titulo}", key=f"tour_ir_{pag}",
                             use_container_width=True, type="primary"):
                    st.session_state.pagina = pag
                    st.session_state.tour_step = 1
                    st.rerun()
                st.markdown("<br style='margin:0'>", unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("← Voltar", use_container_width=False):
                st.session_state.tour_step = 2
                st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
#  PÁGINAS PRINCIPAIS
# ═══════════════════════════════════════════════════════════════════════════════

def pagina_inicio():
    st.markdown('<div style="font-size:2rem; font-weight:800; color:#0F172A; margin-bottom:0.3rem;">Olá, Professora! 👋</div>', unsafe_allow_html=True)
    st.markdown('<div style="color:#64748B; font-size:1rem; margin-bottom:0.5rem;">Escolha o que você quer fazer hoje:</div>', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3)
    _atalhos = [
        ("plano",    "✏️", "Criar Plano de Aula",
         "Digite o tema da aula e o sistema monta o plano com os objetivos BNCC certos",
         "#ECFDF5", "#A7F3D0", "#065F46"),
        ("alunos",   "👶", "Relatório dos Alunos",
         "Registre como cada aluno está se desenvolvendo e gere o relatório bimestral",
         "#EFF6FF", "#BFDBFE", "#1D4ED8"),
        ("profabot", "💬", "Tirar uma Dúvida",
         "Pergunte qualquer coisa sobre a BNCC ou sobre como usar o sistema",
         "#F5F3FF", "#DDD6FE", "#5B21B6"),
    ]
    for col, (pag, icon, titulo, desc, bg, border, cor) in zip([c1, c2, c3], _atalhos):
        with col:
            st.markdown(f"""<div style="background:{bg}; border:1.5px solid {border}; border-radius:14px; padding:1.5rem; box-shadow:0 2px 10px rgba(0,0,0,0.05);">
                <div style="font-size:2rem; margin-bottom:0.5rem;">{icon}</div>
                <div style="font-weight:700; color:{cor}; margin-bottom:0.3rem; font-size:1rem;">{titulo}</div>
                <div style="color:#475569; font-size:0.82rem; line-height:1.5;">{desc}</div>
            </div>""", unsafe_allow_html=True)
            if st.button(titulo, key=f"at_{pag}", use_container_width=True, type="primary"):
                st.session_state.pagina = pag
                st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="sec">💡 Dica do Dia</div>', unsafe_allow_html=True)
    _dicas = [
        "Os campos de experiência da BNCC não são disciplinas — eles integram diferentes saberes de forma lúdica e contextualizada.",
        "Ao planejar uma atividade, pense primeiro na vivência que a criança vai ter, depois o sistema te ajuda a associar os objetivos BNCC.",
        "A mesma atividade pode ter objetivos BNCC diferentes para bebês (EI01) e crianças maiores (EI02). O ProfaPlanner diferencia automaticamente!",
        "Argila e massinha de modelar se encaixam no campo 'Traços, Sons, Cores e Formas' — experimente digitar esse tema no Plano de Aula!",
        "Conflitos entre crianças são oportunidades pedagógicas — a BNCC tem objetivos específicos para isso no campo 'O eu, o outro e o nós'.",
    ]
    st.markdown(f"""<div style="background:#FFFFFF; border:1.5px solid #E2E8F0; border-radius:14px; padding:1.2rem 1.5rem; box-shadow:0 2px 8px rgba(0,0,0,0.04);">
        <div style="color:#334155; font-size:0.93rem; line-height:1.7;">{random.choice(_dicas)}</div>
    </div>""", unsafe_allow_html=True)

    # Link para retornar ao tour
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("🔄  Rever a apresentação do sistema", use_container_width=False):
        st.session_state.pagina = "boas_vindas"
        st.session_state.tour_step = 1
        st.rerun()


# ──────────────────────────────────────────────────────────────────────────────

def pagina_modelo():
    st.markdown('<div class="sec">📋 Meu Modelo de Plano de Aula</div>', unsafe_allow_html=True)
    st.markdown('<p style="color:#94A3B8;">Aqui você coloca o formulário de plano de aula que a sua escola usa. '
                'O sistema usa ele para gerar os documentos no formato certo.</p>', unsafe_allow_html=True)

    col_upload, col_info = st.columns([1.3, 0.7])

    with col_upload:

        # ── Arquivos detectados automaticamente na pasta ──────────────────
        arquivos_locais = detectar_arquivos_locais()
        if arquivos_locais:
            st.markdown("""<div class="card card-verde" style="padding:1rem 1.2rem;">
                <div style="font-weight:700; color:#10B981; margin-bottom:0.5rem;">
                    📁 Encontrei este(s) arquivo(s) na pasta do ProfaPlanner:
                </div>
            </div>""", unsafe_allow_html=True)

            for arq in arquivos_locais:
                c_nome, c_btn = st.columns([3, 1])
                with c_nome:
                    st.markdown(f'<div style="color:#CBD5E1; padding:0.3rem 0;">📄 <strong>{arq.name}</strong></div>',
                                unsafe_allow_html=True)
                with c_btn:
                    if st.button("Usar este", key=f"usar_{arq.name}", type="primary"):
                        ext = arq.suffix.lower()
                        try:
                            if ext in [".docx", ".doc"]:
                                from docx import Document
                                doc = Document(str(arq))
                                paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
                                st.session_state.modelo_escola = {
                                    "nome": arq.name,
                                    "tipo": "Word",
                                    "paragrafos": paragrafos[:15],
                                    "caminho": str(arq),
                                }
                            else:
                                import pandas as pd
                                df = pd.read_excel(str(arq))
                                st.session_state.modelo_escola = {
                                    "nome": arq.name,
                                    "tipo": "Excel",
                                    "colunas": list(df.columns),
                                    "caminho": str(arq),
                                }
                            st.success(f"✅ **{arq.name}** carregado como modelo!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Não consegui ler o arquivo: {e}")

            st.divider()

        # ── Upload manual ─────────────────────────────────────────────────
        st.markdown("**Ou faça o upload de outro arquivo:**")
        arquivo = st.file_uploader(
            "Selecione o arquivo do modelo",
            type=["xlsx", "xls", "docx", "doc"],
            help="Formatos aceitos: Word (.docx) ou Excel (.xlsx)",
        )
        if arquivo:
            ext = Path(arquivo.name).suffix.lower()
            with st.spinner("Lendo o arquivo..."):
                try:
                    if ext in [".docx", ".doc"]:
                        from docx import Document
                        doc = Document(io.BytesIO(arquivo.read()))
                        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
                        st.session_state.modelo_escola = {
                            "nome": arquivo.name, "tipo": "Word", "paragrafos": paragrafos[:15],
                        }
                    else:
                        import pandas as pd
                        df = pd.read_excel(arquivo)
                        st.session_state.modelo_escola = {
                            "nome": arquivo.name, "tipo": "Excel", "colunas": list(df.columns),
                        }
                    st.success(f"✅ **{arquivo.name}** carregado com sucesso!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Não consegui ler o arquivo. Tente novamente. ({e})")

        st.divider()

        # ── Template padrão ───────────────────────────────────────────────
        st.markdown("**Não tem o arquivo agora?**")
        st.info("💡 Sem problema! Você pode usar o modelo padrão do ProfaPlanner. "
                "Depois, quando tiver o documento da escola, pode trocar.")
        if not st.session_state.modelo_escola:
            if st.button("📄  Usar o modelo padrão do sistema", type="primary", use_container_width=True):
                st.session_state.modelo_escola = {"nome": "Modelo Padrão ProfaPlanner", "tipo": "padrao"}
                st.success("✅ Modelo padrão ativado! Você já pode criar planos de aula.")
                st.rerun()
        else:
            if st.button("🔄  Trocar o modelo", use_container_width=True):
                st.session_state.modelo_escola = None
                st.rerun()

    with col_info:
        st.markdown("""<div class="card">
            <div style="font-weight:700; color:#10B981; margin-bottom:1rem;">📌 Para que serve isso?</div>
            <div style="color:#475569; font-size:0.87rem; line-height:2;">
                A sua escola provavelmente tem um formulário específico para o plano de aula.<br><br>
                Ao adicionar esse arquivo aqui, o sistema vai <strong style="color:#1E293B;">gerar os documentos
                no formato que a sua escola pede</strong>, já preenchido com os objetivos BNCC certos. 📋
            </div>
        </div>""", unsafe_allow_html=True)

        _m = st.session_state.modelo_escola
        if _m:
            st.markdown(f"""<div class="card card-verde">
                <div style="font-weight:700; color:#10B981;">✅ Modelo Ativo</div>
                <div style="color:#CBD5E1; margin-top:0.4rem;">📄 {_m.get('nome','—')}</div>
                <div style="color:#475569; font-size:0.78rem;">Tipo: {_m.get('tipo','—')}</div>
            </div>""", unsafe_allow_html=True)


# ──────────────────────────────────────────────────────────────────────────────

def pagina_plano_aula():
    # ── Inicializar estado do fluxo ──────────────────────────────────────────
    if "plano_step" not in st.session_state:
        st.session_state.plano_step = 1
    if "objetivos_selecionados" not in st.session_state:
        st.session_state.objetivos_selecionados = []
    if "plano_gerado" not in st.session_state:
        st.session_state.plano_gerado = None

    step = st.session_state.plano_step

    # ── Cabeçalho ─────────────────────────────────────────────────────────────
    st.markdown("""
    <div style='margin-bottom:1.5rem;'>
        <div class='pp-hero-title'>✨ Criar Sequência Didática</div>
        <div class='pp-hero-sub'>Siga os 3 passos abaixo e a IA escreve o plano completo para você!</div>
    </div>
    """, unsafe_allow_html=True)

    # ── Barra de Progresso em 3 Etapas ────────────────────────────────────────
    s1 = "pp-step-active" if step == 1 else ("pp-step-done" if step > 1 else "pp-step-pending")
    s2 = "pp-step-active" if step == 2 else ("pp-step-done" if step > 2 else "pp-step-pending")
    s3 = "pp-step-active" if step == 3 else ("pp-step-done" if step > 3 else "pp-step-pending")
    n1 = "✓" if step > 1 else "1"
    n2 = "✓" if step > 2 else "2"
    n3 = "✓" if step > 3 else "3"

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(f"""<div class="pp-step {s1}">
            <span class="pp-step-num" style="background:rgba(255,255,255,0.25);">{n1}</span>
            📝 Dados da Aula
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown(f"""<div class="pp-step {s2}">
            <span class="pp-step-num" style="background:rgba(255,255,255,0.25);">{n2}</span>
            🎯 Objetivos BNCC
        </div>""", unsafe_allow_html=True)
    with c3:
        st.markdown(f"""<div class="pp-step {s3}">
            <span class="pp-step-num" style="background:rgba(255,255,255,0.25);">{n3}</span>
            🤖 Plano Gerado
        </div>""", unsafe_allow_html=True)

    st.markdown("<div class='pp-divider'></div>", unsafe_allow_html=True)

    _faixas = {"EI01": "👶 Bebês (0 a 1a6m)", "EI02": "🧒 Crianças (1a7m a 3a11m)", "Ambas": "Mista"}
    _campos = {
        "Todos": "🔍 Todos os campos",
        "EO": "🤝 O eu, o outro e o nós",
        "CG": "🏃 Corpo, gestos e movimentos",
        "TS": "🎨 Traços, sons, cores e formas",
        "EF": "💬 Escuta, fala, pensamento e imaginação",
        "ET": "🔢 Espaços, tempos, quantidades e transformações",
    }

    # ══════════════════════════════════════════════════════════════════════════
    # ETAPA 1 — DADOS DA AULA
    # ══════════════════════════════════════════════════════════════════════════
    if step == 1:
        st.markdown("### Dados da Sequência Didática", unsafe_allow_html=False)

        titulo_sd = st.text_input(
            "Título da Sequência *",
            placeholder="Ex: Descobrindo meu corpo através das brincadeiras",
            help="Dê um nome criativo à sua SD"
        )

        c_seg, c_faixa = st.columns(2)
        with c_seg:
            segmento = st.selectbox(
                "Segmento / Turma",
                ["Berçário 1", "Berçário 2", "Maternal 1", "Maternal 2", "Pré 1", "Pré 2"]
            )
        with c_faixa:
            faixa_sel = st.selectbox("Faixa Etária", options=list(_faixas.keys()), format_func=lambda x: _faixas[x])

        c_dur, c_camp = st.columns(2)
        with c_dur:
            duracao = st.selectbox("Duração da SD", ["Diário", "Semanal", "Quinzenal", "Mensal", "Bimestral"])
        with c_camp:
            campo_sel = st.selectbox("Campo Principal de Experiência", options=list(_campos.keys()), format_func=lambda x: _campos[x])

        tema = st.text_input(
            "Tema / Atividade Central *",
            placeholder="Ex: Brincadeira com argila e água",
            help="O tema principal que será trabalhado na sequência"
        )

        # Guarda na sessão para usar nas próximas etapas
        st.session_state._sd_titulo = titulo_sd
        st.session_state._sd_segmento = segmento
        st.session_state._sd_faixa = faixa_sel
        st.session_state._sd_duracao = duracao
        st.session_state._sd_campo = campo_sel
        st.session_state._sd_tema = tema

        if tema.strip():
            _, col_btn, _ = st.columns([1, 2, 1])
            with col_btn:
                if st.button("Próximo: Encontrar objetivos BNCC →", type="primary", use_container_width=True):
                    st.session_state.plano_step = 2
                    st.rerun()
        else:
            st.markdown("""
            <div class='pp-card-coral' style='padding:1rem;'>
                <span style='color:#EA580C; font-weight:600;'>💡 Dica:</span>
                <span style='color:#9A3412;'> Preencha o "Tema" para avançar para os objetivos BNCC!</span>
            </div>
            """, unsafe_allow_html=True)

    # ══════════════════════════════════════════════════════════════════════════
    # ETAPA 2 — OBJETIVOS BNCC
    # ══════════════════════════════════════════════════════════════════════════
    elif step == 2:
        tema = st.session_state.get("_sd_tema", "")
        faixa_sel = st.session_state.get("_sd_faixa", "EI01")
        campo_sel = st.session_state.get("_sd_campo", "Todos")

        st.markdown(f"""
        <div class='pp-card-verde' style='padding:1rem 1.4rem;'>
            <span style='color:#059669; font-weight:600;'>🎯 Tema selecionado:</span>
            <span style='color:#065F46; font-size:1.05rem; font-weight:700;'> {tema}</span>
        </div>
        """, unsafe_allow_html=True)

        faixa_filtro = None if faixa_sel == "Ambas" else faixa_sel
        campo_filtro = None if campo_sel == "Todos" else campo_sel

        with st.spinner("🔍 Buscando os objetivos BNCC mais relevantes..."):
            resultados = engine.buscar(texto=tema, faixa=faixa_filtro, campo=campo_filtro)[:5]

        if resultados:
            st.markdown("""
            <div class='pp-section-title'>
                🎯 Objetivos BNCC encontrados — marque os que se aplicam à sua SD:
            </div>
            """, unsafe_allow_html=True)
            st.markdown("<div style='color:#64748B; font-size:0.85rem; margin-bottom:0.8rem;'>Todos já vêm marcados. Desmarque se não quiser incluir algum.</div>", unsafe_allow_html=True)

            selecionados = []
            for obj in resultados:
                campo_nome = _campos.get(obj.get('campo', ''), obj.get('campo', ''))
                col_chk, col_info = st.columns([0.05, 0.95])
                with col_chk:
                    checked = st.checkbox("", value=True, key=f"obj_{obj['codigo']}", label_visibility="collapsed")
                with col_info:
                    st.markdown(f"""
                    <div class="pp-obj-card" style="{'border-color:#10B981; background:#ECFDF5;' if True else ''}">
                        <div style='display:flex; align-items:center; gap:8px; margin-bottom:4px;'>
                            <span style='background:#10B981; color:white; padding:2px 8px; border-radius:6px; font-size:0.78rem; font-weight:700;'>{obj['codigo']}</span>
                            <span class='pp-tag pp-tag-verde'>{campo_nome}</span>
                        </div>
                        <div style='color:#1E293B; font-size:0.88rem; line-height:1.5;'>{obj['descricao']}</div>
                    </div>
                    """, unsafe_allow_html=True)
                if checked:
                    selecionados.append(obj)

            st.session_state.objetivos_selecionados = selecionados

            st.markdown("<div class='pp-divider'></div>", unsafe_allow_html=True)

            c_back, c_next = st.columns(2)
            with c_back:
                if st.button("← Voltar e editar dados", use_container_width=True):
                    st.session_state.plano_step = 1
                    st.rerun()
            with c_next:
                if selecionados:
                    if st.button(f"🤖 Gerar Plano Completo com IA →", type="primary", use_container_width=True):
                        with st.spinner("✨ A IA está escrevendo sua Sequência Didática completa... isso pode levar uns 15 segundos!"):
                            from modules.chatbot import ProfaBot
                            bot = ProfaBot()
                            plano_gerado = bot.gerar_plano_completo(
                                tema=tema,
                                faixa=_faixas[faixa_sel],
                                campo=_campos.get(campo_sel, campo_sel),
                                objetivos=selecionados
                            )
                            if plano_gerado:
                                st.session_state.plano_gerado = plano_gerado
                                st.session_state.plano_step = 3
                                st.rerun()
                            else:
                                st.error("Houve um erro ao gerar o plano. Tente novamente.")
                else:
                    st.warning("Selecione pelo menos 1 objetivo para continuar.")
        else:
            st.warning("Não encontrei objetivos para essas palavras. Volte e tente um tema diferente.")
            if st.button("← Voltar"):
                st.session_state.plano_step = 1
                st.rerun()

    # ══════════════════════════════════════════════════════════════════════════
    # ETAPA 3 — PLANO GERADO (REVISÃO E DOWNLOAD)
    # ══════════════════════════════════════════════════════════════════════════
    elif step == 3:
        pg = st.session_state.plano_gerado
        tema = st.session_state.get("_sd_tema", "")
        faixa_sel = st.session_state.get("_sd_faixa", "EI01")
        campo_sel = st.session_state.get("_sd_campo", "Todos")
        segmento = st.session_state.get("_sd_segmento", "")
        duracao = st.session_state.get("_sd_duracao", "Diário")
        titulo_sd = st.session_state.get("_sd_titulo", "Sequência Didática")

        st.markdown("""
        <div class='pp-card-verde' style='padding:1.2rem 1.5rem;'>
            <div style='display:flex; align-items:center; gap:10px;'>
                <span style='font-size:1.8rem;'>🎉</span>
                <div>
                    <div style='font-weight:700; color:#065F46; font-size:1.05rem;'>Plano criado com sucesso!</div>
                    <div style='color:#059669; font-size:0.88rem;'>Revise abaixo e baixe quando estiver pronto.</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Resumo rápido
        col_j, col_o = st.columns(2)
        with col_j:
            st.markdown("<div class='pp-section-title'>📖 Justificativa</div>", unsafe_allow_html=True)
            justificativa = st.text_area("Justificativa", value=pg.get("justificativa", ""), height=120, label_visibility="collapsed")
        with col_o:
            st.markdown("<div class='pp-section-title'>🎯 Objetivo Geral</div>", unsafe_allow_html=True)
            obj_geral = st.text_area("Objetivo Geral", value=pg.get("obj_geral", ""), height=120, label_visibility="collapsed")

        st.markdown("<div class='pp-divider'></div>", unsafe_allow_html=True)
        st.markdown("<div class='pp-section-title'>🧩 Atividades Propostas</div>", unsafe_allow_html=True)
        st.markdown("<div style='color:#64748B; font-size:0.85rem; margin-bottom:0.8rem;'>Clique em cada atividade para expandir e editar.</div>", unsafe_allow_html=True)

        atividades = []
        for i, ativ in enumerate(pg.get("atividades", [])):
            titulo_expander = f"Atividade {i+1}: {ativ.get('nome', '')}"
            with st.expander(titulo_expander, expanded=False):
                nome_a = st.text_input("Nome da Atividade", value=ativ.get("nome", ""), key=f"n_{i}")
                desc_a = st.text_area("Passo a Passo Detalhado", value=ativ.get("descricao", ""), height=150, key=f"d_{i}")
                atividades.append({"nome": nome_a, "descricao": desc_a})

        st.markdown("<div class='pp-divider'></div>", unsafe_allow_html=True)
        st.markdown("<div class='pp-section-title'>📊 Avaliação e Registro</div>", unsafe_allow_html=True)
        avaliacao = st.text_area("Como o aprendizado será observado e registrado", value=pg.get("avaliacao", ""), height=100, label_visibility="collapsed")

        # Feedback para IA
        st.markdown("""
        <div class='pp-card-azul' style='padding:1.2rem; margin-top:1rem;'>
            <div style='font-weight:700; color:#1D4ED8; margin-bottom:0.4rem;'>🤖 Quer ajustar algo?</div>
            <div style='color:#1E40AF; font-size:0.88rem;'>Diga à IA o que deve mudar e ela reescreve o plano para você!</div>
        </div>
        """, unsafe_allow_html=True)
        feedback_ia = st.text_input("O que você quer mudar?", placeholder="Ex: As atividades estão muito complexas para bebês, inclua mais músicas e movimentos suaves.")

        c_refazer, _ = st.columns([1, 2])
        with c_refazer:
            if feedback_ia and st.button("🔄 Refazer com IA", type="primary", use_container_width=True):
                with st.spinner("Reescrevendo plano..."):
                    from modules.chatbot import ProfaBot
                    bot = ProfaBot()
                    novo_plano = bot.gerar_plano_completo(
                        tema=tema,
                        faixa=_faixas[faixa_sel],
                        campo=_campos.get(campo_sel, campo_sel),
                        objetivos=st.session_state.objetivos_selecionados,
                        feedbacks=feedback_ia
                    )
                    if novo_plano:
                        st.session_state.plano_gerado = novo_plano
                        st.rerun()

        st.markdown("<div class='pp-divider'></div>", unsafe_allow_html=True)
        st.markdown("""
        <div style='font-size:1.05rem; font-weight:700; color:#0F172A; margin-bottom:0.8rem;'>
            ⬇️ Tudo certo? Baixe seu documento!
        </div>
        """, unsafe_allow_html=True)

        plano_final = {
            "titulo": titulo_sd if titulo_sd else "Sequência Didática",
            "tema": tema,
            "segmento": segmento,
            "faixa": _faixas[faixa_sel],
            "campo": _campos.get(campo_sel, campo_sel),
            "duracao": duracao,
            "justificativa": justificativa,
            "obj_geral": obj_geral,
            "avaliacao": avaliacao,
            "atividades": atividades,
            "objetivos_bncc": st.session_state.objetivos_selecionados,
        }

        c_word, c_pdf = st.columns(2)
        from modules.exportador import Exportador
        exportador = Exportador()

        with c_word:
            doc_word = exportador.gerar_sd_word(plano_final)
            nome_word = f"SD_{tema[:15].replace(' ', '_').lower()}.docx"
            st.download_button(
                "📄 Baixar em Word (.docx)",
                data=doc_word, file_name=nome_word,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with c_pdf:
            doc_pdf = exportador.gerar_sd_pdf(plano_final)
            nome_pdf = f"SD_{tema[:15].replace(' ', '_').lower()}.pdf"
            st.download_button(
                "📕 Baixar em PDF",
                data=doc_pdf, file_name=nome_pdf,
                mime="application/pdf",
                use_container_width=True
            )

        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🆕 Criar nova Sequência Didática", use_container_width=False):
            for k in ["plano_step", "plano_gerado", "objetivos_selecionados", "_sd_tema", "_sd_titulo",
                      "_sd_faixa", "_sd_campo", "_sd_segmento", "_sd_duracao", "plano_atual"]:
                if k in st.session_state:
                    del st.session_state[k]
            st.rerun()



def pagina_alunos():
    st.markdown('<div class="sec">👶 Relatório dos Alunos</div>', unsafe_allow_html=True)
    
    from modules.db import BancoDados
    db = BancoDados()
    alunos_db = db.get_alunos()

    tab1, tab2, tab3 = st.tabs(["📋  Minha Turma", "📊  Registro de Desenvolvimento", "💾 Backup / Restauração"])

    with tab1:
        st.markdown("**Adicionar aluno à turma:**")
        c1, c2, c3 = st.columns([2.5, 2, 0.8])
        with c1:
            nome_aluno = st.text_input("Nome da criança", placeholder="Nome completo",
                                       label_visibility="collapsed")
        with c2:
            faixa_aluno = st.selectbox("Faixa",
                                       ["EI01 — Bebês (0 a 1a 6m)", "EI02 — Crianças maiores (1a7m a 3a11m)"],
                                       label_visibility="collapsed")
        with c3:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("➕ Adicionar", type="primary", use_container_width=True):
                if nome_aluno.strip():
                    faixa_cod = "EI01" if "EI01" in faixa_aluno else "EI02"
                    nomes = [a["nome"].lower() for a in alunos_db]
                    if nome_aluno.strip().lower() in nomes:
                        st.error("Já existe um aluno com esse nome.")
                    else:
                        db.add_aluno(nome_aluno.strip(), faixa_cod)
                        st.success(f"✅ {nome_aluno} adicionado(a)!")
                        st.rerun()
                else:
                    st.error("Digite o nome da criança.")

        st.divider()

        if alunos_db:
            st.markdown(f"**{len(alunos_db)} criança(s) na turma:**")
            for aluno in alunos_db:
                c_nome, c_prog, c_del = st.columns([3, 2, 0.4])
                with c_nome:
                    tag = "tag-ei01" if aluno["faixa"] == "EI01" else "tag-ei02"
                    label = "Bebê" if aluno["faixa"] == "EI01" else "Criança maior"
                    st.markdown(f'<span style="color:#1E293B; font-weight:500;">👤 {aluno["nome"]}</span>'
                                f'&nbsp;&nbsp;<span class="{tag}">{label}</span>',
                                unsafe_allow_html=True)
                with c_prog:
                    n_av = len(aluno.get("evolucao", {}))
                    n_at = sum(1 for v in aluno.get("evolucao", {}).values() if v == "A")
                    st.markdown(f'<span style="color:#475569; font-size:0.82rem;">'
                                f'{n_av} avaliados · {n_at} atingidos</span>', unsafe_allow_html=True)
                with c_del:
                    if st.button("🗑️", key=f"del_{aluno['id']}"):
                        db.delete_aluno(aluno["id"])
                        st.rerun()
                st.divider()
        else:
            st.info("Ainda não há crianças cadastradas. Adicione os alunos da sua turma acima.")

    with tab2:
        if not alunos_db:
            st.info("Cadastre os alunos na aba **Minha Turma** primeiro.")
        else:
            aluno_nome = st.selectbox("Qual criança você quer avaliar?",
                                      [a["nome"] for a in alunos_db])
            aluno = next(a for a in alunos_db if a["nome"] == aluno_nome)

            _status_opts = {"A": "✅ Já consegue fazer", "D": "🔄 Está aprendendo", "N": "⏳ Ainda não iniciou"}

            st.markdown(
                f'<div style="color:#64748B; font-size:0.85rem; margin:0.5rem 0 1rem;">'
                f"Para cada habilidade abaixo, marque como <strong style='color:#1E293B;'>{aluno_nome}</strong> está:</div>",
                unsafe_allow_html=True)

            for campo_nome, objetivos in engine.get_objetivos_por_faixa(aluno["faixa"]).items():
                n_at = sum(1 for o in objetivos if aluno["evolucao"].get(o["codigo"]) == "A")
                with st.expander(f"{campo_nome} — {n_at}/{len(objetivos)} conquistados", icon="📌"):
                    for obj in objetivos:
                        c_desc, c_st = st.columns([4, 1.4])
                        with c_desc:
                            st.markdown(
                                f'<span class="badge-bncc">{obj["codigo"]}</span>'
                                f'<span style="color:#334155; font-size:0.85rem; margin-left:6px;">{obj["descricao"]}</span>',
                                unsafe_allow_html=True)
                        with c_st:
                            status = aluno["evolucao"].get(obj["codigo"], "N")
                            novo = st.selectbox("s", list(_status_opts.keys()),
                                                format_func=lambda x: _status_opts[x],
                                                index=list(_status_opts.keys()).index(status),
                                                key=f"st_{aluno_nome}_{obj['codigo']}",
                                                label_visibility="collapsed")
                            
                            if novo != status:
                                evolucao_atual = aluno["evolucao"].copy()
                                evolucao_atual[obj["codigo"]] = novo
                                db.update_evolucao(aluno["id"], evolucao_atual)
                                st.rerun()
                        st.divider()

            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("⬇️  Gerar Relatório Bimestral desta criança",
                         type="primary", use_container_width=True):
                with st.spinner(f"Gerando relatório de {aluno_nome}..."):
                    doc_bytes = Exportador().gerar_relatorio_aluno(aluno, engine)
                nome_arquivo = f"relatorio_{aluno_nome.replace(' ', '_').lower()}.docx"
                st.download_button(
                    label=f"📄  Baixar Relatório de {aluno_nome}",
                    data=doc_bytes, file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)

    with tab3:
        st.markdown('<div style="font-weight:700; color:#0F172A; font-size:1.1rem; margin-bottom:0.5rem;">💾 Backup / Restauração da Turma Inteira</div>', unsafe_allow_html=True)
        st.markdown('<div style="color:#475569; font-size:0.9rem; margin-bottom:1rem;">Use esta aba para baixar os dados da sua turma atual ou restaurar uma turma salva anteriormente. Ideal para manter seus dados seguros ou se estiver usando outro computador.</div>', unsafe_allow_html=True)
        
        c_down, c_up = st.columns(2)
        with c_down:
            with st.container(border=True):
                st.markdown('<div style="color:#1E293B; font-weight:700; margin-bottom:0.3rem;">⬇️ Exportar (Backup)</div>', unsafe_allow_html=True)
                st.markdown('<div style="color:#475569; font-size:0.85rem; margin-bottom:1rem;">Baixe a planilha com todos os alunos e avaliações. Guarde este arquivo em segurança.</div>', unsafe_allow_html=True)
                try:
                    excel_bytes = db.exportar_alunos_excel()
                    st.download_button(
                        label="Baixar Turma para Excel",
                        data=excel_bytes,
                        file_name="Meus_Alunos_ProfaPlanner.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Erro ao gerar backup: {e}")
                
        with c_up:
            with st.container(border=True):
                st.markdown('<div style="color:#1E293B; font-weight:700; margin-bottom:0.3rem;">⬆️ Importar (Restauração)</div>', unsafe_allow_html=True)
                st.markdown('<div style="color:#475569; font-size:0.85rem; margin-bottom:0.5rem;">Faça o upload de uma planilha de backup baixada anteriormente.</div>', unsafe_allow_html=True)
                uploaded_file = st.file_uploader("Arquivo Excel", type=["xlsx"], label_visibility="collapsed")
                if uploaded_file is not None:
                    if st.button("Restaurar Turma", type="primary", use_container_width=True):
                        try:
                            db.importar_alunos_excel(uploaded_file.read())
                            st.success("✅ Turma restaurada com sucesso!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erro ao importar. Certifique-se que o arquivo é válido. Detalhes: {e}")

# ──────────────────────────────────────────────────────────────────────────────

def pagina_profabot():
    st.markdown('''
    <div style="margin-bottom:1rem;">
        <div style="font-size:1.5rem; font-weight:800; color:#0F172A; margin-bottom:0.3rem;">💬 Tirar Dúvidas com a Assistente Virtual</div>
        <div style="color:#475569; font-size:0.95rem;">Pode perguntar à vontade sobre a BNCC, sobre planejamento ou sobre como usar o sistema. Ela responde em linguagem simples! 😊</div>
    </div>
    ''', unsafe_allow_html=True)

    _sugestoes = [
        "Qual código BNCC usar com argila?",
        "Como escrever o relatório de um bebê?",
        "O que é o campo 'Traços, sons, cores e formas'?",
        "Diferença entre EI01 e EI02",
        "Atividades para trabalhar o campo EO",
        "Como usar o sistema para criar um plano?",
    ]
    st.markdown("**💡 Perguntas frequentes — clique para enviar:**")
    _cols = st.columns(2)
    for _i, _s in enumerate(_sugestoes):
        with _cols[_i % 2]:
            if st.button(f"💬 {_s}", key=f"sug_{_i}", use_container_width=True):
                st.session_state.historico_chat.append({"role": "user", "content": _s})
                with st.spinner("Pensando..."):
                    bot = ProfaBot(st.session_state.gemini_key)
                    resp = bot.responder(_s, st.session_state.historico_chat[:-1])
                st.session_state.historico_chat.append({"role": "assistant", "content": resp})
                st.rerun()

    st.divider()

    for msg in st.session_state.historico_chat:
        bg = "#EFF6FF" if msg["role"] == "user" else "#F0FDF4"
        border = "#BFDBFE" if msg["role"] == "user" else "#A7F3D0"
        avatar = "👩‍🏫" if msg["role"] == "user" else "🤖"
        st.markdown(
            f"""<div style="background:{bg};border:1px solid {border};border-radius:12px;
            padding:0.9rem 1.1rem;margin-bottom:0.5rem;display:flex;gap:10px;align-items:flex-start;">
                <span style="font-size:1.4rem;flex-shrink:0;">{avatar}</span>
                <div style="color:#1E293B;font-size:0.93rem;line-height:1.6;">{msg["content"]}</div>
            </div>""",
            unsafe_allow_html=True
        )

    st.markdown("<br>", unsafe_allow_html=True)
    c_inp, c_btn = st.columns([5, 1])
    with c_inp:
        prompt = st.text_input("Sua pergunta", placeholder="Digite sua pergunta aqui...",
                               label_visibility="collapsed", key="chat_input_field")
    with c_btn:
        enviar = st.button("Enviar", type="primary", use_container_width=True)

    if enviar and prompt:
        st.session_state.historico_chat.append({"role": "user", "content": prompt})
        with st.spinner("Pensando..."):
            bot = ProfaBot(st.session_state.gemini_key)
            resp = bot.responder(prompt, st.session_state.historico_chat[:-1])
        st.session_state.historico_chat.append({"role": "assistant", "content": resp})
        st.rerun()

    if st.session_state.historico_chat:
        if st.button("Limpar conversa"):
            st.session_state.historico_chat = []
            st.rerun()



# ═══════════════════════════════════════════════════════════════════════════════
#  ROTEAMENTO
# ═══════════════════════════════════════════════════════════════════════════════
_PAGINAS = {
    "boas_vindas": pagina_boas_vindas,
    "inicio":      pagina_inicio,
    "modelo":      pagina_modelo,
    "plano":       pagina_plano_aula,
    "alunos":      pagina_alunos,
    "profabot":    pagina_profabot,
}
_PAGINAS[st.session_state.pagina]()
